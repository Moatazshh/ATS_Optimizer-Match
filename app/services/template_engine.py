from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from app.models.schemas import StructuredResume, CoverLetter
import io

# ─────────────────────────────────────────────────────────────────────────────
# RULE: Use ONLY plain Normal-style paragraphs + runs.
# Never use List Bullet, tables, text boxes, or headers/footers.
# Never put critical info in document headers/footers.
# Bullets = "• " prefix on a Normal paragraph with a small indent.
# This approach consistently achieves 95-100% ATS parse rates.
# ─────────────────────────────────────────────────────────────────────────────

class ATSDocxGenerator:
    """Generates an ATS-optimised DOCX resume using only plain-text structure."""

    def __init__(self, resume: StructuredResume, template_type: str = "general"):
        self.resume = resume
        self.template_type = template_type.lower()
        self.doc = Document()
        self._setup_document()

    # ── Document setup ────────────────────────────────────────────────────────
    def _setup_document(self):
        # 1-inch margins — the universal ATS-safe default
        for section in self.doc.sections:
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin   = Inches(1.0)
            section.right_margin  = Inches(1.0)

        # Override only the Normal style font — no XML manipulation
        normal = self.doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        # Metadata (helps ATS classify the file)
        self.doc.core_properties.title    = f"{self.resume.contact_info.name} Resume"
        self.doc.core_properties.author   = "ATS Optimizer"
        self.doc.core_properties.language = "en-US"

    # ── Low-level paragraph helpers ───────────────────────────────────────────
    def _para(self, text: str = "", bold: bool = False, size: int = 11,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_before: int = 0,
              space_after: int = 4) -> None:
        """Add a plain Normal paragraph."""
        p = self.doc.add_paragraph(style="Normal")
        p.alignment = align
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
        return p

    def _section_header(self, label: str) -> None:
        """Bold all-caps section header with visible line separator."""
        # Blank spacing paragraph
        sp = self.doc.add_paragraph(style="Normal")
        sp.paragraph_format.space_before = Pt(10)
        sp.paragraph_format.space_after  = Pt(0)

        p = self.doc.add_paragraph(style="Normal")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(label.upper())
        run.bold = True
        run.font.size = Pt(12)

        # Horizontal rule via a series of underscores (pure text, 100% parseable)
        hr = self.doc.add_paragraph(style="Normal")
        hr.paragraph_format.space_before = Pt(0)
        hr.paragraph_format.space_after  = Pt(6)
        hr.add_run("─" * 80)  # Em-dash line — purely decorative text

    def _bullet(self, text: str) -> None:
        """Add a bullet point as a plain Normal paragraph (no List style)."""
        p = self.doc.add_paragraph(style="Normal")
        p.paragraph_format.space_before   = Pt(0)
        p.paragraph_format.space_after    = Pt(2)
        p.paragraph_format.left_indent    = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.add_run("• " + text)

    # ── Document assembly ─────────────────────────────────────────────────────
    def generate(self) -> io.BytesIO:
        self._add_contact_info()

        if self.template_type in ("tech",):
            self._add_summary(); self._add_skills()
            self._add_experience(); self._add_education()
            self._add_certifications()
        elif self.template_type == "finance":
            self._add_summary(); self._add_experience()
            self._add_education(); self._add_skills()
            self._add_certifications()
        elif self.template_type == "executive":
            self._add_summary(); self._add_experience()
            self._add_skills(); self._add_education()
        else:   # general / healthcare
            self._add_summary(); self._add_skills()
            self._add_experience(); self._add_education()
            self._add_certifications()

        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf

    # ── Section builders ──────────────────────────────────────────────────────
    def _add_contact_info(self):
        c = self.resume.contact_info

        # Name — bold, large, centred
        p_name = self.doc.add_paragraph(style="Normal")
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_after = Pt(4)
        run_name = p_name.add_run(c.name)
        run_name.bold = True
        run_name.font.size = Pt(16)

        # Contact line — all on one paragraph, centred
        parts = []
        if c.email:    parts.append(c.email)
        if c.phone:    parts.append(c.phone)
        if c.location: parts.append(c.location)
        if c.linkedin: parts.append(c.linkedin)
        if c.website:  parts.append(c.website)

        p_contact = self.doc.add_paragraph(style="Normal")
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_contact.paragraph_format.space_after = Pt(6)
        p_contact.add_run(" | ".join(parts))

    def _add_summary(self):
        if not self.resume.summary:
            return
        self._section_header("Summary")
        p = self.doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(self.resume.summary)

    def _add_skills(self):
        if not self.resume.skills:
            return
        self._section_header("Skills")
        # Comma-separated — every ATS recognises this format
        p = self.doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(", ".join(self.resume.skills))

    def _add_experience(self):
        if not self.resume.experience:
            return
        self._section_header("Work Experience")
        for exp in self.resume.experience:
            # Title | Company
            p_title = self.doc.add_paragraph(style="Normal")
            p_title.paragraph_format.space_before = Pt(6)
            p_title.paragraph_format.space_after  = Pt(0)
            r_title = p_title.add_run(f"{exp.title}")
            r_title.bold = True
            p_title.add_run(f"  |  {exp.company}")

            # Dates | Location
            p_date = self.doc.add_paragraph(style="Normal")
            p_date.paragraph_format.space_before = Pt(0)
            p_date.paragraph_format.space_after  = Pt(2)
            date_str = f"{exp.start_date} – {exp.end_date}"
            if exp.location:
                date_str += f"  |  {exp.location}"
            p_date.add_run(date_str).italic = True

            # Bullets
            for bp in exp.bullet_points:
                self._bullet(bp)

    def _add_education(self):
        if not self.resume.education:
            return
        self._section_header("Education")
        for edu in self.resume.education:
            p = self.doc.add_paragraph(style="Normal")
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(0)
            p.add_run(f"{edu.degree}").bold = True
            p.add_run(f",  {edu.institution}")

            p2 = self.doc.add_paragraph(style="Normal")
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(4)
            loc = f"  |  {edu.location}" if edu.location else ""
            p2.add_run(f"{edu.graduation_date}{loc}").italic = True

    def _add_certifications(self):
        if not self.resume.certifications:
            return
        self._section_header("Certifications")
        for cert in self.resume.certifications:
            self._bullet(cert)


# ─────────────────────────────────────────────────────────────────────────────
# PDF Generator  (kept as-is — it was working)
# ─────────────────────────────────────────────────────────────────────────────
class ATSPDFGenerator(FPDF):
    def __init__(self, resume: StructuredResume, template_type: str = "general"):
        super().__init__()
        self.resume = resume
        self.template_type = template_type.lower()
        self.set_auto_page_break(auto=True, margin=15)
        self.add_page()
        self.l_margin = 15
        self.r_margin = 15
        self.set_left_margin(15)
        self.set_right_margin(15)
        self.set_font("Arial", size=10)
        self.set_title(f"{self.resume.contact_info.name} Resume")
        self.set_author("ATS Optimizer")
        self.set_subject("Professional Resume optimized for ATS systems")

    def _clean_text(self, text: str) -> str:
        if not text: return ""
        replacements = {
            "\u2013": "-",  "\u2014": "-",
            "\u2018": "'",  "\u2019": "'",
            "\u201c": '"',  "\u201d": '"',
            "\u2022": "-",  "\u2026": "...",
            "\u2500": "-",  # horizontal box-drawing char
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        return text.encode("latin-1", "ignore").decode("latin-1")

    def _add_section_header(self, text: str):
        self.ln(5)
        self.set_x(self.l_margin)
        self.set_font("Arial", "B", 12)
        self.cell(0, 8, text.upper(), ln=1, align="L")
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(2)
        self.set_font("Arial", size=10)

    def generate(self) -> io.BytesIO:
        self.set_y(15)
        self.set_font("Arial", "B", 16)
        self.cell(0, 10, self._clean_text(self.resume.contact_info.name), ln=1, align="C")
        self.set_font("Arial", size=10)

        parts = []
        if self.resume.contact_info.email:    parts.append(self.resume.contact_info.email)
        if self.resume.contact_info.phone:    parts.append(self.resume.contact_info.phone)
        if self.resume.contact_info.location: parts.append(self.resume.contact_info.location)
        if self.resume.contact_info.linkedin: parts.append(self.resume.contact_info.linkedin)
        if self.resume.contact_info.website:  parts.append(self.resume.contact_info.website)
        self.cell(0, 5, self._clean_text(" | ".join(parts)), ln=1, align="C")

        self._add_summary()
        if self.template_type == "tech":
            self._add_skills(); self._add_experience()
        else:
            self._add_experience(); self._add_skills()
        self._add_education()
        self._add_certifications()

        buf = io.BytesIO()
        self.output(buf)
        buf.seek(0)
        return buf

    def _add_summary(self):
        if not self.resume.summary: return
        self._add_section_header("Summary")
        self.set_x(self.l_margin)
        self.multi_cell(0, 5, self._clean_text(self.resume.summary))

    def _add_skills(self):
        if not self.resume.skills: return
        self._add_section_header("Skills")
        self.set_x(self.l_margin)
        self.multi_cell(0, 5, self._clean_text(", ".join(self.resume.skills)))

    def _add_experience(self):
        if not self.resume.experience: return
        self._add_section_header("Work Experience")
        for exp in self.resume.experience:
            self.set_x(self.l_margin)
            self.set_font("Arial", "B", 10)
            self.cell(0, 5, self._clean_text(f"{exp.title} | {exp.company}"), ln=1)

            self.set_x(self.l_margin)
            self.set_font("Arial", "I", 9)
            loc_str = f" | {exp.location}" if exp.location else ""
            self.cell(0, 5, self._clean_text(f"{exp.start_date} - {exp.end_date}{loc_str}"), ln=1)

            self.set_font("Arial", size=10)
            for bp in exp.bullet_points:
                self.set_x(self.l_margin + 5)
                avail = self.w - self.l_margin - self.r_margin - 5
                self.multi_cell(avail, 5, self._clean_text(f"* {bp}"))
            self.ln(2)

    def _add_education(self):
        if not self.resume.education: return
        self._add_section_header("Education")
        for edu in self.resume.education:
            self.set_font("Arial", "B", 10)
            self.cell(0, 5, self._clean_text(f"{edu.degree}, {edu.institution}"), ln=True)
            self.set_font("Arial", size=9)
            loc_str = f" | {edu.location}" if edu.location else ""
            self.cell(0, 5, self._clean_text(f"{edu.graduation_date}{loc_str}"), ln=True)

    def _add_certifications(self):
        if not self.resume.certifications: return
        self._add_section_header("Certifications")
        self.set_font("Arial", size=10)
        for cert in self.resume.certifications:
            self.set_x(self.l_margin + 5)
            self.cell(0, 5, self._clean_text(f"* {cert}"), ln=True)


# ─────────────────────────────────────────────────────────────────────────────
# Public helpers
# ─────────────────────────────────────────────────────────────────────────────
def generate_docx(resume: StructuredResume, template_type: str = "general") -> io.BytesIO:
    return ATSDocxGenerator(resume, template_type).generate()

def generate_pdf(resume: StructuredResume, template_type: str = "general") -> io.BytesIO:
    return ATSPDFGenerator(resume, template_type).generate()

def generate_cover_letter_docx(cl: CoverLetter, user_name: str = "Candidate Name") -> io.BytesIO:
    doc = Document()

    p = doc.add_paragraph(cl.date)
    p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(cl.recipient_name)
    p = doc.add_paragraph(cl.company_name)
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph(f"{cl.salutation} {cl.recipient_name},")
    p.paragraph_format.space_after = Pt(12)

    for para in cl.content.split("\n\n"):
        clean = para.strip()
        if not clean: continue
        if clean.lower().startswith("dear") or clean.lower().startswith("sincerely"):
            continue
        p = doc.add_paragraph(clean)
        p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph(cl.closing + ",")
    doc.add_paragraph()
    doc.add_paragraph(user_name)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generate_cover_letter_pdf(cl: CoverLetter, user_name: str = "Candidate Name") -> io.BytesIO:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=11)

    pdf.cell(0, 10, cl.date, ln=True)
    pdf.ln(5)
    pdf.cell(0, 5, cl.recipient_name, ln=True)
    pdf.cell(0, 5, cl.company_name, ln=True)
    pdf.ln(10)
    pdf.cell(0, 10, f"{cl.salutation} {cl.recipient_name},", ln=True)
    pdf.ln(5)

    for para in cl.content.split("\n\n"):
        clean = para.strip()
        if not clean: continue
        if clean.lower().startswith("dear") or clean.lower().startswith("sincerely"):
            continue
        encoded = clean.encode("latin-1", "ignore").decode("latin-1")
        pdf.multi_cell(0, 6, encoded, align="L")
        pdf.ln(5)

    pdf.ln(5)
    pdf.cell(0, 5, cl.closing + ",", ln=True)
    pdf.ln(5)
    pdf.cell(0, 5, user_name, ln=True)

    buf = io.BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf
